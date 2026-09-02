"""DOCX generator using python-docx per TRD Section 5 and Security Document Section 6."""

import logging
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.resume import (
    AchievementEntry, CertificationEntry, EducationEntry,
    ExperienceEntry, PersonalDetails, ProjectEntry, Resume,
    SkillGroup, ReferencesMode, ReferenceEntry,
)

logger = logging.getLogger(__name__)

TEMPLATE_COLORS = {
    "modern": RGBColor(0x1a, 0x1a, 0x2e),
    "classic": RGBColor(0x2c, 0x3e, 0x50),
    "compact": RGBColor(0x1b, 0x1b, 0x1b),
}


class DOCXGenerator:
    def generate(self, resume: Resume, template_id: str = "modern") -> bytes:
        try:
            doc = Document()
            color = TEMPLATE_COLORS.get(template_id, TEMPLATE_COLORS["modern"])
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(10)

            self._add_personal_details(doc, resume.personal_details, color)
            if resume.professional_summary:
                self._add_heading(doc, "Professional Summary", color)
                doc.add_paragraph(resume.professional_summary)
            if resume.experience:
                self._add_heading(doc, "Experience", color)
                self._add_experience(doc, resume.experience)
            if resume.education:
                self._add_heading(doc, "Education", color)
                self._add_education(doc, resume.education)
            if resume.projects:
                self._add_heading(doc, "Projects", color)
                self._add_projects(doc, resume.projects)
            if resume.skills:
                self._add_heading(doc, "Skills", color)
                self._add_skills(doc, resume.skills)
            if resume.certifications:
                self._add_heading(doc, "Certifications", color)
                self._add_certifications(doc, resume.certifications)
            if resume.achievements:
                self._add_heading(doc, "Achievements", color)
                self._add_achievements(doc, resume.achievements)
            if resume.references:
                self._add_heading(doc, "References", color)
                self._add_references(doc, resume.references)

            buf = BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()
            logger.info("docx_generated", extra={"detail": f"{len(docx_bytes)} bytes"})
            return docx_bytes
        except Exception as e:
            logger.error("docx_generation_error", extra={"detail": str(e)[:200]})
            raise

    def _add_heading(self, doc: Document, text: str, color: RGBColor):
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = color

    def _add_personal_details(self, doc: Document, pd: PersonalDetails, color: RGBColor):
        title = doc.add_heading(pd.full_name, level=0)
        for run in title.runs:
            run.font.color.rgb = color
            run.font.size = Pt(20)
        contact_parts = []
        if pd.professional_title:
            contact_parts.append(pd.professional_title)
        if pd.email:
            contact_parts.append(pd.email)
        if pd.phone:
            contact_parts.append(pd.phone)
        if pd.location:
            contact_parts.append(pd.location)
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

    def _add_experience(self, doc: Document, entries: List[ExperienceEntry]):
        for entry in entries:
            p = doc.add_paragraph()
            run = p.add_run(f"{entry.job_title} at {entry.company_name}")
            run.bold = True
            doc.add_paragraph(f"{entry.start_date} - {entry.end_date}")
            if entry.description_bullets:
                for bullet in entry.description_bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

    def _add_education(self, doc: Document, entries: List[EducationEntry]):
        for entry in entries:
            p = doc.add_paragraph()
            run = p.add_run(f"{entry.degree} - {entry.institution_name}")
            run.bold = True
            if entry.end_date:
                doc.add_paragraph(entry.end_date)
            if entry.details:
                doc.add_paragraph(entry.details)

    def _add_projects(self, doc: Document, entries: List[ProjectEntry]):
        for entry in entries:
            p = doc.add_paragraph()
            run = p.add_run(entry.project_name)
            run.bold = True
            if entry.description_bullets:
                for bullet in entry.description_bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

    def _add_skills(self, doc: Document, groups: List[SkillGroup]):
        for group in groups:
            label = f"{group.category_label}: " if group.category_label else ""
            doc.add_paragraph(f"{label}{', '.join(group.skills)}")

    def _add_certifications(self, doc: Document, entries: List[CertificationEntry]):
        for entry in entries:
            text = entry.certification_name
            if entry.issuing_organization:
                text += f" - {entry.issuing_organization}"
            doc.add_paragraph(text)

    def _add_achievements(self, doc: Document, entries: List[AchievementEntry]):
        for entry in entries:
            doc.add_paragraph(entry.statement, style="List Bullet")

    def _add_references(self, doc: Document, refs):
        if isinstance(refs, ReferencesMode):
            doc.add_paragraph("Available upon request")
        elif isinstance(refs, list):
            for ref in refs:
                doc.add_paragraph(f"{ref.name} - {ref.relationship or ''}")


docx_generator = DOCXGenerator()
