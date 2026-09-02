"""DOCX parser integration using python-docx per TRD Section 5."""

import logging
from io import BytesIO

from docx import Document

logger = logging.getLogger(__name__)


class DOCXParser:
    def extract_text(self, file_bytes: bytes) -> str:
        try:
            doc = Document(BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    style_name = (para.style.name or "").lower() if para.style else ""
                    prefix = ""
                    if "heading" in style_name:
                        prefix = "\n## "
                    elif "list" in style_name:
                        prefix = "- "
                    text_parts.append(f"{prefix}{para.text.strip()}")

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            full_text = "\n".join(text_parts)
            logger.info("docx_parsed", extra={"detail": f"Extracted {len(full_text)} chars"})
            return full_text
        except Exception as e:
            logger.error("docx_parse_error", extra={"detail": str(e)[:200]})
            raise


docx_parser = DOCXParser()
